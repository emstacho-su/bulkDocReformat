# writer.py
#
# Revised to:
#   - Keep existing logic for “X.” → Heading 1 (title), “X.X” → Heading 2, “X.X.X” → Heading 3
#   - Preserve all numeric prefixes in each node["heading"]
#   - Apply run.font.size = Pt(11) to every body‐text run
#   - Do NOT touch headers or footers (we leave Policy Reference for later)

from pathlib import Path
import re
from typing import Dict, Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def write_new_doc(parsed: Dict[str, Any], output_path: Path) -> None:
    """
    Given the parsed structure (from parser.py), produce a brand-new .docx at output_path:
      - Use Heading 1 for the document title
      - Use Heading 2 for each top‐level section (parsed["sections"] entries)
      - Use Heading 3 for each second‐level (node["children"])
      - Use Heading 4 for each third‐level (grandchildren), etc.
      - Every run of “body” text (node["content"]) is explicitly set to 11 pt.
      - We do NOT insert any footer or header here.
    """

    # Start a brand‐new document (Normal.dotm defaults apply).
    doc = Document()

    # 1) Title (assume the very first heading is the document title).
    title = parsed.get("document_title", "").strip()
    if title:
        title_para = doc.add_paragraph(title, style="Heading 1")
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Override size/bold per corporate spec:
        for run in title_para.runs:
            run.font.size = Pt(16)
            run.font.bold = True

    # 2) Recursively write sections
    def _write_node(node: Dict[str, Any], level: int) -> None:
        """
        Recursively write one node and its children.
          level == 1 → Heading 2
          level == 2 → Heading 3
          level == 3 → Heading 4
          etc.
        """
        heading_text = node.get("heading", "").strip()
        if not heading_text:
            return

        # Choose the correct style name based on `level`:
        if level == 1:
            style_name = "Heading 2"
        elif level == 2:
            style_name = "Heading 3"
        else:
            style_name = "Heading 4"

        # Write the heading with numeric prefix preserved
        h_para = doc.add_paragraph(heading_text, style=style_name)
        # Override font size/bold for heading runs if needed
        for run in h_para.runs:
            if style_name == "Heading 2":
                run.font.size = Pt(14)
                run.font.bold = True
            elif style_name == "Heading 3":
                run.font.size = Pt(12)
                run.font.bold = True
            # Heading 4 and deeper inherit default; you can override here if desired.

        # 2.a) Write any “content” lines under this heading
        for line in node.get("content", []):
            if line.strip():
                body_para = doc.add_paragraph()
                run = body_para.add_run(line.strip())
                run.font.size = Pt(11)

        # 2.b) Recurse into children (each child increments `level` by 1)
        for child in node.get("children", []):
            _write_node(child, level + 1)

    # Kick off with level=1 for each top‐level section
    for top in parsed.get("sections", []):
        _write_node(top, level=1)

    # 3) Save the result
    doc.save(str(output_path))
