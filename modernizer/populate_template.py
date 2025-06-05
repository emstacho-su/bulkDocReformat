from pathlib import Path
from typing import Dict, Any, List
import re
import logging

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt   # add next to your other imports
from math import ceil
from docx.shared import Pt
# Import the parser
from modernizer.parser import parse_legacy_docx_by_sequence

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


def _strip_all_numbers(s: str) -> str:
    return re.sub(r"^\s*\d+(?:\.\d+)*\s*", "", s).strip()


def _find_paragraph(doc: Document, exact: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    logger.warning(f"Placeholder paragraph '{exact}' not found in template.")
    return None


def _remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = None


def _insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str = None, indent_level: int = 0) -> Paragraph:
    text = re.sub(r"\s+", " ", text).strip()
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    if style:
        new_para.style = style
    if text:
        new_para.text = text

    if indent_level:
        new_para.paragraph_format.left_indent = Inches(0. * indent_level)

    return new_para


def _insert_definition_after(paragraph: Paragraph, full_text: str) -> Paragraph:
    full_text = re.sub(r"\s+", " ", full_text).strip()
    if ":" in full_text:
        term, rest = full_text.split(":", 1)
        term = term.strip() + ":"
        rest = rest.strip()
    else:
        term = full_text
        rest = ""

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    run1 = new_para.add_run(term)
    run1.bold = True
    if rest:
        new_para.add_run(" " + rest)

    return new_para


def normalize(text: str) -> str:
    return text.strip().lower()


def strip_numeric_prefix(text: str) -> str:
    m = re.match(r"^[\d\.]+\s*(.*)$", text.strip())
    return (m.group(1).strip() if m else text.strip())


def center_align_cell(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _find_revision_table(doc: Document) -> Table:
    def normalize_hdr(s):
        return re.sub(r"[^\w]", "", s.strip().lower())

    for tbl in doc.tables:
        if len(tbl.rows) >= 1 and len(tbl.columns) >= 3:
            headers = [normalize_hdr(cell.text) for cell in tbl.rows[0].cells[:3]]
            if ("rev" in headers[0] and "date" in headers[1] and ("desc" in headers[2] or "description" in headers[2])):
                return tbl
    return None


def _get_legacy_revision_column_indices(header_row: List[str]) -> Dict[str, int]:
    col_map = {}
    for idx, val in enumerate(header_row):
        val_l = val.lower()
        if "chg" in val_l and "no" in val_l:
            col_map["rev"] = idx
        elif "date" in val_l:
            col_map["date"] = idx
        elif "desc" in val_l or "change" in val_l:
            col_map["desc"] = idx
    return col_map


def insert_revision_history_table(doc: Document, rev_data: Dict[str, Any]) -> None:
    if rev_data.get("type") != "table":
        logger.info("Revision history is not a table; skipping for now.")
        return

    table = _find_revision_table(doc)
    if not table:
        logger.warning("Could not find [Rev, Date, Description] table in template.")
        return

    legacy_rows = rev_data["rows"]
    if not legacy_rows or len(legacy_rows) < 2:
        logger.warning("Legacy revision table has no data rows.")
        return

    header_row = legacy_rows[0]
    col_map = _get_legacy_revision_column_indices(header_row)

    missing = [k for k in ["rev", "date", "desc"] if k not in col_map]
    if missing:
        logger.warning(f"Missing expected legacy revision columns: {missing}")
        return

    for row in legacy_rows[1:]:
        if len(row) <= max(col_map.values()):
            continue
        new_row = table.add_row()
        for p in new_row.cells[0].paragraphs + new_row.cells[1].paragraphs + new_row.cells[2].paragraphs:
            p.clear()
        new_row.cells[0].text = row[col_map["rev"]].strip()
        new_row.cells[1].text = row[col_map["date"]].strip()
        new_row.cells[2].text = row[col_map["desc"]].strip()
        center_align_cell(new_row.cells[0])
        center_align_cell(new_row.cells[1])


# ----------------------------------------------------------------------
#  Pagination helper – shrink the Revision-History table *iff* the
#  combined intro block barely overruns one page (cell-level counting).
# ----------------------------------------------------------------------
def _fit_intro_sections_to_page(doc: Document) -> None:
    """
    Count “lines” (estimating 80 chars per paragraph‐line, 40 chars per cell‐line)
    from the first paragraph up to “Definitions”.  Shrink the Revision-History
    table if total lines are between 47 and 70.
    """
    # 1) Find the first 'Definitions' paragraph
    def_idx = next(
        (i for i, p in enumerate(doc.paragraphs)
         if p.text.strip().lower().startswith("definitions")), None
    )
    if def_idx is None:
        return

    # 2) Estimate lines: paragraphs first (80 chars/line), tables next
    lines = 0
    seen_tables = set()

    for p in doc.paragraphs[:def_idx]:
        if p._element.getparent().tag.endswith("tc"):
            seen_tables.add(p._tc.table)
            continue

        txt = p.text.strip()
        if txt:
            lines += max(1, ceil(len(txt) / 80))

    # --- START OF REPLACED BLOCK: count per cell instead of per row ---
    for tbl in seen_tables:
        for row in tbl.rows:
            for cell in row.cells:
                cell_txt = cell.text.strip()
                if not cell_txt:
                    continue
                # Assume ~40 chars wrap per visible line inside a table cell
                lines += max(1, ceil(len(cell_txt) / 40))
    # --- END OF REPLACED BLOCK ---

    # 3) Decide: ≤46 → fits; 47–70 → shrink; >70 → leave alone
    if lines <= 38 or lines > 70:
        return

    # 4) Shrink only the Revision-History table
    rev_tbl = _find_revision_table(doc)
    if rev_tbl is None:
        return

    for row in rev_tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)


def populate_template(
    parsed: Dict[str, Any],
    template_path: Path,
    output_path: Path
    ) -> None:


    """
    Open the template, then fill in each section based on the parsed structure.
    1) PURPOSE + SCOPE combined → fill under "Purpose and Scope"
    2) Remove the placeholders "PURPOSE:" and "SCOPE:" entirely
    3) REVISION HISTORY
    4) PROCESS OWNER / PROCESS DESIGNEES
    5) DEFINITIONS
    6) PROCEDURES  
    7) REFERENCES (including "Related Documents" sub‐headings)
    8) RECORDS     
    9) POLICY REFERENCE (skip/remove if not present)
    """
    doc = Document(str(template_path))

    # ------------------------------------------------------------------
    # 1) PURPOSE + SCOPE combined → under "Purpose and Scope"
    # ------------------------------------------------------------------
    p_header = _find_paragraph(doc, "Purpose and Scope")
    if p_header and parsed.get("purpose_scope_block"):
        last_p = p_header
        for line in parsed["purpose_scope_block"].splitlines():
            if line.strip():
                last_p = _insert_paragraph_after(last_p, text=line)
    else:
        if not p_header:
            logger.warning("No 'Purpose and Scope' placeholder found; skipping insertion.")
        if not parsed.get("purpose_scope_block"):
            logger.info("No Purpose/Scope block parsed; leaving template blank under that heading.")

    # After inserting, remove any standalone placeholders
    to_remove = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt in {"PURPOSE:", "Purpose:", "SCOPE:", "Scope"}:
            to_remove.append(p)
    for p in to_remove:
        _remove_paragraph(p)

     # ------------------------------------------------------------------
    # 2) REVISION HISTORY
    # ------------------------------------------------------------------
    insert_revision_history_table(doc, parsed.get("revision_history", {}))
    # ------------------------------------------------------------------
    # 3) PROCESS OWNER / PROCESS DESIGNEES  (populate existing 2×2 with borders)
    # ------------------------------------------------------------------
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # 1) Gather all owners and designees from parsed["sections"]
    owners = []
    designees = []
    for sec in parsed["sections"]:
        head = normalize(strip_numeric_prefix(sec["heading"]))
        if head.startswith("process owner"):
            owners = [c["heading"].strip() for c in sec.get("children", [])]
        elif head.startswith("process designees"):
            designees = [c["heading"].strip() for c in sec.get("children", [])]

    owner_str = ", ".join(owners)
    designee_str = ", ".join(designees)

    # Log what will go into each cell
    logger.info(f"Process Owner/Authorized By: will be filled with: {owner_str}")
    logger.info(f"Process Designees: will be filled with: {designee_str}")

    # 2) Find any existing 2×2 table placeholder
    placeholder_tbl = None
    for tbl in doc.tables:
        # Check if exactly 2 rows, and first row has exactly 2 cells
        if len(tbl.rows) == 2 and len(tbl.rows[0].cells) == 2:
            placeholder_tbl = tbl
            break

    # 3) If none found, warn and create a new 2×2 at document end
    if placeholder_tbl is None:
        logger.warning("No existing 2×2 table found. Creating a new one at document end.")
        placeholder_tbl = doc.add_table(rows=2, cols=2)

    # 4) Clear any existing text in that 2×2
    for r in range(2):
        for c in range(2):
            placeholder_tbl.cell(r, c).text = ""

    # 5) A helper to add borders to a given cell
    def set_cell_border(cell):
        """
        Add a single-line border on all four sides of this cell.
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove existing borders if any
        for child in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(child)

        tcBorders = OxmlElement("w:tcBorders")

        # Border attributes: size=4 (width), val=single, space=0, color=000000 (black)
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)

        tcPr.append(tcBorders)

    # 6) Populate & format each cell (bold + center-align), then add borders
    # Row 0, Col 0: “Process Owner/Authorized By: ”
    cell00 = placeholder_tbl.cell(0, 0)
    p00 = cell00.paragraphs[0]
    run00 = p00.add_run("Process Owner/Authorized By: ")
    run00.bold = True
    p00.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell00.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell00)

    # Row 0, Col 1: comma-separated owner names
    cell01 = placeholder_tbl.cell(0, 1)
    p01 = cell01.paragraphs[0]
    run01 = p01.add_run(owner_str)
    run01.bold = True
    p01.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell01.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell01)

    # Row 1, Col 0: “Process Designees: ”
    cell10 = placeholder_tbl.cell(1, 0)
    p10 = cell10.paragraphs[0]
    run10 = p10.add_run("Process Designees: ")
    run10.bold = True
    p10.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell10.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell10)

    # Row 1, Col 1: comma-separated designee names
    cell11 = placeholder_tbl.cell(1, 1)
    p11 = cell11.paragraphs[0]
    run11 = p11.add_run(designee_str)
    run11.bold = True
    p11.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell11.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell11)

    # ------------------------------------------------------------------
    # -- Condense intro to one page if feasible ------------------------
    # ------------------------------------------------------------------
    _fit_intro_sections_to_page(doc)

    # ------------------------------------------------------------------
    # 4) DEFINITIONS
    # ------------------------------------------------------------------
    p_def_heading = _find_paragraph(doc, "Definitions")
    defs_node = next(
        (sec for sec in parsed["sections"]
         if "definitions" in normalize(strip_numeric_prefix(sec["heading"]))),
        None
    )
    if p_def_heading and defs_node:
        last_para = p_def_heading
        for child in defs_node.get("children", []):
            full_text = child["heading"]
            if full_text.strip():
                last_para = _insert_definition_after(last_para, full_text=full_text)
                for cont in child.get("content", []):
                    for subline in cont.splitlines():
                        if subline.strip():
                            last_para = _insert_paragraph_after(last_para, text=subline)
    else:
        if not p_def_heading:
            logger.warning("No 'Definitions' placeholder found; skipping definitions.")
        if defs_node is None:
            logger.info("No 'Definitions' section in parsed doc; leaving placeholder blank.")


    # ------------------------------------------------------------------
    # 5) PROCEDURES  (4.x → Heading 2 auto‐numbered as “2.x”)
    # ------------------------------------------------------------------
    p_proc_heading = _find_paragraph(doc, "Procedures")
    proc_node = next(
        (sec for sec in parsed["sections"]
         if "procedures" in normalize(strip_numeric_prefix(sec["heading"]))),
        None
    )
    if p_proc_heading and proc_node:
        last_para = p_proc_heading
        for child in proc_node.get("children", []):
            raw = re.sub(r"\s+", " ", child["heading"]).strip()
            just_text = re.sub(r"^\s*4\.\d+\s*", "", raw)
            last_para = _insert_paragraph_after(
                last_para,
                text=just_text,
                style="Heading 2"
            )
            # Insert each body line under Heading 2 with indent_level=1
            for line in child.get("content", []):
                for subline in line.splitlines():
                    if subline.strip():
                        last_para = _insert_paragraph_after(
                            last_para,
                            text=subline,
                            indent_level=1  # ← NEW
                        )
            for subsub in child.get("children", []):
                raw_ss = re.sub(r"\s+", " ", subsub["heading"]).strip()
                just_ss = re.sub(r"^\s*4\.\d+\.(\d+)\s*", r"\1 ", raw_ss)
                last_para = _insert_paragraph_after(
                    last_para,
                    text=just_ss,
                    style="Heading 3"
                )
                # Insert each body line under Heading 3 with indent_level=2
                for line2 in subsub.get("content", []):
                    for subline2 in line2.splitlines():
                        if subline2.strip():
                            last_para = _insert_paragraph_after(
                                last_para,
                                text=subline2,
                                indent_level=2  # ← NEW
                            )
    else:
        if not p_proc_heading:
            logger.warning("No 'Procedures' placeholder found; skipping procedures.")
        if proc_node is None:
            logger.info("No 'Procedures' section in parsed doc; leaving placeholder blank.")

    # ------------------------------------------------------------------
    # 6) REFERENCES  (5.x → Heading 2 auto‐numbered as “3.x”)
    #    Also insert any "Related Documents" children as sub‐headings here.
    # ------------------------------------------------------------------
    p_ref_heading = _find_paragraph(doc, "References")
    ref_node = next(
        (sec for sec in parsed["sections"]
         if "references" in normalize(strip_numeric_prefix(sec["heading"]))),
        None
    )
    related_node = next(
        (sec for sec in parsed["sections"]
         if "related documents" in normalize(strip_numeric_prefix(sec["heading"]))),
        None
    )

    if p_ref_heading:
        last_para = p_ref_heading

        # Insert “References” children
        if ref_node:
            for child in ref_node.get("children", []):
                raw = re.sub(r"\s+", " ", child["heading"]).strip()
                just_text = re.sub(r"^\s*5\.\d+\s*", "", raw)
                last_para = _insert_paragraph_after(
                    last_para,
                    text=just_text,
                    style="Heading 2"
                )
                # Insert each body line under Heading 2 with indent_level=1
                for line in child.get("content", []):
                    for subline in line.splitlines():
                        if subline.strip():
                            last_para = _insert_paragraph_after(
                                last_para,
                                text=subline,
                                indent_level=1  # ← NEW
                            )
            # Insert any loose lines under "References" with indent_level=1
            for line in ref_node.get("content", []):
                for subline in line.splitlines():
                    if subline.strip():
                        last_para = _insert_paragraph_after(
                            last_para,
                            text=subline,
                            indent_level=1  # ← NEW
                        )
        else:
            logger.info("No 'References' section in parsed doc; leaving placeholder blank.")

        # Insert “Related Documents” if present
        if related_node:
            for child in related_node.get("children", []):
                raw = re.sub(r"\s+", " ", child["heading"]).strip()
                last_para = _insert_paragraph_after(
                    last_para,
                    text=raw,
                    style="Heading 2"
                )
                for line in child.get("content", []):
                    for subline in line.splitlines():
                        if subline.strip():
                            last_para = _insert_paragraph_after(
                                last_para,
                                text=subline,
                                indent_level=1  # ← NEW
                            )
    else:
        logger.warning("No 'References' placeholder found; skipping references/related documents.")

    # ------------------------------------------------------------------
    # 7) RECORDS  (6.x → Heading 2 auto‐numbered as “4.x”)
    # ------------------------------------------------------------------
    p_rec = _find_paragraph(doc, "Records")
    records_node = next(
        (s for s in parsed["sections"]
         if normalize(strip_numeric_prefix(s["heading"])).startswith("records")),
        None
    )
    if p_rec and records_node:
        last = p_rec
        # Child headings first
        for child in records_node["children"]:
            txt = _strip_all_numbers(child["heading"])
            last = _insert_paragraph_after(
                last,
                text=txt,
                style="Heading 2"
            )
            # Indent each child line under Heading 2 by one level
            for ln in child.get("content", []):
                if ln.strip():
                    last = _insert_paragraph_after(
                        last,
                        text=ln,
                        indent_level=1  # ← NEW
                    )
        # Then any loose content lines under "Records" with indent_level=1
        for ln in records_node.get("content", []):
            if ln.strip():
                last = _insert_paragraph_after(
                    last,
                    text=ln,
                    indent_level=1  # ← NEW
                )
    else:
        if not p_rec:
            logger.warning("No 'Records' placeholder found; skipping records.")
        if records_node is None:
            logger.info("No 'Records' section in parsed doc; leaving placeholder blank.")

                  # ------------------------------------------------------------------
    # 8) POLICY REFERENCE  (7.x → Heading 2 auto‑numbered as “5.x”)
    #    We flatten every descendant into plain indented lines.
    # ------------------------------------------------------------------
    pol_node = next(
        (sec for sec in parsed["sections"]
         if "policy reference" in normalize(strip_numeric_prefix(sec["heading"]))),
        None
    )
    p_pol_heading = _find_paragraph(doc, "Policy Reference")

    if p_pol_heading and pol_node:
        last_para = p_pol_heading

        flat_lines = []

        # -- helper to walk one level of descendants -------------------
        def collect(node):
            """Append node heading + its content lines to flat_lines."""
            head_txt = _strip_all_numbers(node["heading"]).strip()
            if head_txt:
                flat_lines.append(head_txt)

            # direct content
            for ln in node.get("content", []):
                for subln in ln.splitlines():
                    if subln.strip():
                        flat_lines.append(subln.strip())

            # one more level (grand‑children) — rare but safe
            for gchild in node.get("children", []):
                g_head = _strip_all_numbers(gchild["heading"]).strip()
                if g_head:
                    flat_lines.append(g_head)
                for ln2 in gchild.get("content", []):
                    for subln2 in ln2.splitlines():
                        if subln2.strip():
                            flat_lines.append(subln2.strip())

        # collect every child (and its descendants)
        for child in pol_node.get("children", []):
            collect(child)

        # loose lines on the Policy‑Reference node itself
        for ln in pol_node.get("content", []):
            for subln in ln.splitlines():
                if subln.strip():
                    flat_lines.append(subln.strip())

        # insert them in order, indented 0.25"
        for txt in flat_lines:
            last_para = _insert_paragraph_after(
                last_para,
                text=txt,
                indent_level=1
            )

    else:
        # Placeholder missing OR legacy doc had no Policy Reference
        if p_pol_heading and not pol_node:
            logger.info("No 'Policy Reference' in parsed doc; removing placeholder.")
            _remove_paragraph(p_pol_heading)
        elif pol_node and not p_pol_heading:
            logger.warning(
                "Parsed a 'Policy Reference' section but template has no placeholder. "
                "Skipping Policy Reference insertion."
            )




    # ------------------------------------------------------------------
    # 9) SAVE (with "-copy" logic if the file already exists)
    # ------------------------------------------------------------------
    if output_path.exists():
        base = output_path.stem
        ext = output_path.suffix
        parent = output_path.parent
        counter = 1
        candidate = parent / f"{base}-copy{ext}"
        while candidate.exists():
            counter += 1
            candidate = parent / f"{base}-copy{counter}{ext}"
        output_path = candidate
        logger.info(f"Output file exists; saving as '{output_path.name}' instead.")

    doc.save(str(output_path))
    logger.info(f"Saved modernized document to: {output_path}")