from pathlib import Path
from typing import List, Dict, Any
import re
import logging

from docx import Document

# ---------------------------------------------------------------------
#  Logging
# ---------------------------------------------------------------------
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------
#  Ordered top‑level keywords (lower‑case, stripped of numeric prefixes)
# ---------------------------------------------------------------------
TOP_LEVEL_SEQUENCE = [
    "purpose",          # handled specially
    "scope",
    "definitions",
    "process owner",
    "procedures",
    "references",
    "related documents",
    "records",
    "policy reference",
    "revisions",
]

# ---------------------------------------------------------------------
#  Regex helpers
# ---------------------------------------------------------------------
NUM_PREFIX_PATTERN    = re.compile(r"^[\d\.]+\s*(.*)$")
SUBCLAUSE_PATTERN     = re.compile(r"^\s*\d+\.\d+")
SUBSUBCLAUSE_PATTERN  = re.compile(r"^\s*\d+\.\d+\.\d+")
SINGLE_LEVEL_PATTERN  = re.compile(r"^(\d+)\.\s+")
PURPOSE_AND_SCOPE_PATTERN = re.compile(r"^purpose\s+and\s+scope", re.IGNORECASE)
PROCESS_DESIGNEE_PATTERN  = re.compile(r"^process\s+designee", re.IGNORECASE)
LETTERED_PATTERN      = re.compile(r"^[A-Za-z]\.\s+")
# ---------------------------------------------------------------------
#  Utility functions
# ---------------------------------------------------------------------
def normalize(text: str) -> str:
    return text.strip().lower()

def strip_numeric_prefix(text: str) -> str:
    m = NUM_PREFIX_PATTERN.match(text.strip())
    return m.group(1).strip() if m else text.strip()

def extract_revision_history(doc: Document) -> Dict[str, Any]:
    """Return {"type":"table",...} if last element is table, else written."""
    if doc.tables:
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in doc.tables[-1].rows
        ]
        return {"type": "table", "rows": rows}
    return {"type": "written", "content": []}

# ---------------------------------------------------------------------
#  Main parser
# ---------------------------------------------------------------------
def parse_legacy_docx_by_sequence(path: Path) -> Dict[str, Any]:
    """
    Parse a legacy .docx file and return a structure for populate_template.
    """

    # A)  Load Word file
    try:
        doc = Document(path)
    except Exception as e:
        logger.error(f"Failed to open {path.name}: {e}")
        raise

    # B)  Collect non‑blank paragraphs
    paras = [p for p in doc.paragraphs if p.text.strip()]
    if not paras:
        logger.warning(f"No paragraphs found in {path.name}")
        return {
            "document_title": "",
            "purpose_scope_block": "",
            "sections": [],
            "revision_history": {"type": "written", "content": []},
        }

    # C)  Document title = first bold paragraph (fallback to first line)
    document_title = next(
        (p.text.strip() for p in paras if any(r.bold for r in p.runs)), paras[0].text.strip()
    )
    if document_title == paras[0].text.strip():
        logger.warning(f"No bold title in {path.name}; using first line")

    logger.info(f"[{path.name}] Document Title: {document_title!r}")

    # D)  Locate “Purpose and Scope” and “Definitions”
    idx_ps, idx_def = None, None
    for idx, p in enumerate(paras):
        stripped = strip_numeric_prefix(p.text)
        low = normalize(stripped)
        if idx_ps is None and PURPOSE_AND_SCOPE_PATTERN.match(stripped):
            idx_ps = idx
        if idx_def is None and low.startswith("definitions"):
            idx_def = idx
            break

    if idx_ps is None or idx_def is None or idx_def <= idx_ps:
        logger.warning("Couldn't find both 'Purpose and Scope' and 'Definitions'; "
                       "skipping Purpose/Scope extraction.")
        purpose_scope_block = ""
        idx_def = idx_def if idx_def is not None else len(paras)
    else:
        purpose_scope_block = "\n".join(
            p.text.strip() for p in paras[idx_ps + 1 : idx_def] if p.text.strip()
        )

    # E)  Iterate from Definitions onward, building nested section structure
    remaining = paras[idx_def:]
    sections: List[Dict[str, Any]] = []

    current_top = current_sub = current_subsub = None
    next_top_idx = 2  # 'purpose' and 'scope' consumed logically

    expect_owner_child = False
    capturing_owner_block = False
    owner_buffer: List[str] = []
    capturing_designee_block = False
    designee_buffer: List[str] = []


    def create_top(heading: str):
        nonlocal sections, current_top, current_sub, current_subsub
        nonlocal next_top_idx, expect_owner_child
        nonlocal capturing_owner_block, owner_buffer, capturing_designee_block, designee_buffer

        node = {"heading": heading, "content": [], "children": []}
        sections.append(node)
        current_top, current_sub, current_subsub = node, None, None

        # Reset any in‐progress owner/designee capture whenever we start a fresh top
        expect_owner_child = False
        capturing_owner_block = False
        owner_buffer = []
        capturing_designee_block = False
        designee_buffer = []

        stripped_head = normalize(strip_numeric_prefix(heading))
        if next_top_idx < len(TOP_LEVEL_SEQUENCE) and stripped_head.startswith(
            TOP_LEVEL_SEQUENCE[next_top_idx]
        ):
            next_top_idx += 1


    for p in remaining:
        text = p.text.strip()
        if not text:
            continue

        stripped = strip_numeric_prefix(text)
        low = normalize(stripped)
        raw = text
        is_bold = any(run.bold for run in p.runs)

        # ------------------------------------------------------------------
        #  Special case: promote *any* "Records" heading to a new top
        #  – ignores bold state and punctuation, but still protects against
        #    accidental matches inside body text by requiring either
        #    (a) bold‑face OR (b) a single‑level numeric prefix like "6."+      # ------------------------------------------------------------------        if (
        if(
            low.startswith("records")                # "records", "records:", "records –"
            and (is_bold or SINGLE_LEVEL_PATTERN.match(raw))
        ):
            create_top(text)
            continue
        
        # ------------------------------------------------------------------
        #  Special case: promote *any* "Policy Reference" heading to a new top
        #  – ignores bold state and punctuation, but still protects against
        #    accidental matches inside body text by requiring either
        #    (a) bold-face OR (b) a single-level numeric prefix like "7."
        # ------------------------------------------------------------------
        if (
            low.startswith("policy reference")       # “Policy Reference”, “Policy Reference:”, etc.
            and (is_bold or SINGLE_LEVEL_PATTERN.match(raw))
        ):
            create_top(text)
            continue
        # ------------------------------------------------------------------
        #  Special case: skip any "Revision History" heading or line entirely
        #  – prevents “Revision History …” from becoming content under any section
        # ------------------------------------------------------------------
        if low.startswith("revision history"):
            continue

        # ------------------------------------------------------------------
        #  Handle ongoing Process‑Designee multi‑line block
        # ------------------------------------------------------------------
        if capturing_designee_block:
            is_new_top = (
                current_sub is None
                and SINGLE_LEVEL_PATTERN.match(raw)
                and not SUBCLAUSE_PATTERN.match(raw)
            ) or (
                next_top_idx < len(TOP_LEVEL_SEQUENCE)
                and low.startswith(TOP_LEVEL_SEQUENCE[next_top_idx])
            )
            if is_new_top:
                for line in designee_buffer:
                    token = line.strip()
                    if token:
                        current_top["children"].append(
                            {"heading": token, "content": [], "children": []}
                        )
                designee_buffer = []
                capturing_designee_block = False
            else:
                designee_buffer.append(text)
                continue

        # ------------------------------------------------------------------
        #  Single-level numeric prefix → new top (including “Process Owner”)
        # ------------------------------------------------------------------
        if SINGLE_LEVEL_PATTERN.match(raw) and not SUBCLAUSE_PATTERN.match(raw):
            stripped_head = normalize(strip_numeric_prefix(text))
            create_top(text)

            # If this is “Process Owner”, begin capturing all following owner lines
            if stripped_head.startswith("process owner"):
                capturing_owner_block = True
                owner_buffer = []
            else:
                # Any other top resets both owner/designee state
                capturing_owner_block = False
                owner_buffer = []
                capturing_designee_block = False
                designee_buffer = []

            continue

        # ------------------------------------------------------------------
        #  If we’re in the middle of collecting Process Owner lines
        # ------------------------------------------------------------------
        if capturing_owner_block:
            # Determine if this paragraph starts a new top (e.g. “4. Procedures”) or “Process Designees:”
            is_new_top = (
                SINGLE_LEVEL_PATTERN.match(raw) and not SUBCLAUSE_PATTERN.match(raw)
            ) or (
                next_top_idx < len(TOP_LEVEL_SEQUENCE)
                and low.startswith(TOP_LEVEL_SEQUENCE[next_top_idx])
            )

            # Explicitly treat “Process Designees:” as a top‐breaker, too
            if PROCESS_DESIGNEE_PATTERN.match(stripped):
                is_new_top = True

            if is_new_top:
                # Flush owner_buffer into child nodes under the “Process Owner” top
                for line in owner_buffer:
                    token = line.strip()
                    if token:
                        current_top["children"].append({
                            "heading": token,
                            "content": [],
                            "children": []
                        })
                owner_buffer = []
                capturing_owner_block = False
                # Do NOT continue here, so that this same paragraph can be re‐evaluated
                # (it may start “Process Designees:” or be a fresh top)
            else:
                # Still within “3. Process Owner” block → collect this line
                owner_buffer.append(text)
                continue

        # ------------------------------------------------------------------
        #  Keyword in sequence → new top
        # ------------------------------------------------------------------
        if next_top_idx < len(TOP_LEVEL_SEQUENCE):
            expected = TOP_LEVEL_SEQUENCE[next_top_idx]
            if low.startswith(expected):
                create_top(text)
                if expected == "process owner":
                    expect_owner_child = True
                continue

        # ------------------------------------------------------------------
        #  Process‑Owner child (single line)
        # ------------------------------------------------------------------
        if expect_owner_child:
            child = {"heading": text, "content": [], "children": []}
            current_top["children"].append(child)
            current_sub, current_subsub = child, None
            expect_owner_child = False
            continue

               # ------------------------------------------------------------------
        #  Process-Designee start (“Process Designees:” bold line)
        # ------------------------------------------------------------------
        if PROCESS_DESIGNEE_PATTERN.match(stripped):
            create_top("Process Designees")
            capturing_designee_block = True
            designee_buffer = []

            # If there’s text after the colon on the same line, capture it immediately
            if ":" in stripped:
                remainder = stripped.split(":", 1)[1].strip()
                if remainder:
                    designee_buffer.append(remainder)
            continue


        # ------------------------------------------------------------------
        #  DEFINITIONS: every bold line w/ colon is new child
        # ------------------------------------------------------------------
        if (
            current_top
            and normalize(strip_numeric_prefix(current_top["heading"])).startswith("definitions")
        ):
            is_new_top = (
                current_sub is None
                and SINGLE_LEVEL_PATTERN.match(raw)
                and not SUBCLAUSE_PATTERN.match(raw)
            ) or (
                next_top_idx < len(TOP_LEVEL_SEQUENCE)
                and low.startswith(TOP_LEVEL_SEQUENCE[next_top_idx])
            )
            if not is_new_top:
                if ":" in text:
                    child = {"heading": text, "content": [], "children": []}
                    current_top["children"].append(child)
                    current_sub, current_subsub = child, None
                else:
                    if current_sub:
                        current_sub["heading"] += " " + text
                    else:
                        child = {"heading": text, "content": [], "children": []}
                        current_top["children"].append(child)
                        current_sub = child
                continue

        # ------------------------------------------------------------------
        #  POLICY REFERENCE: every bold line ⇒ new child; plain lines attach
        # ------------------------------------------------------------------
        if (
            current_top
            and normalize(strip_numeric_prefix(current_top["heading"])).startswith("policy reference")
        ):
            is_new_top = (
                current_sub is None
                and SINGLE_LEVEL_PATTERN.match(raw)
                and not SUBCLAUSE_PATTERN.match(raw)
            ) or (
                next_top_idx < len(TOP_LEVEL_SEQUENCE)
                and low.startswith(TOP_LEVEL_SEQUENCE[next_top_idx])
            )
            if not is_new_top:
                if is_bold:
                    child = {"heading": text, "content": [], "children": []}
                    current_top["children"].append(child)
                    current_sub = child
                else:
                    if current_sub:
                        current_sub["content"].append(text)
                    else:
                        child = {"heading": text, "content": [text], "children": []}
                        current_top["children"].append(child)
                        current_sub = child
                continue

        # ------------------------------------------------------------------
        #  Sub‑sub‑clause x.x.x (bold)
        # ------------------------------------------------------------------
        if is_bold and SUBSUBCLAUSE_PATTERN.match(raw) and current_sub:
            child = {"heading": text, "content": [], "children": []}
            current_sub["children"].append(child)
            current_subsub = child
            continue

        # ------------------------------------------------------------------
        #  Sub‑clause x.x (bold) – treat as content under Records
        # ------------------------------------------------------------------
        if is_bold and SUBCLAUSE_PATTERN.match(raw):
            if current_top and normalize(strip_numeric_prefix(current_top["heading"])) == "records":
                current_top["content"].append(text)
                continue
            if current_top:
                child = {"heading": text, "content": [], "children": []}
                current_top["children"].append(child)
                current_sub, current_subsub = child, None
            continue

        # ------------------------------------------------------------------
        #  Other bold line
        # ------------------------------------------------------------------
        if is_bold:
            if current_top and normalize(strip_numeric_prefix(current_top["heading"])) == "records":
                current_top["content"].append(text)
                continue
            if current_subsub:
                current_subsub["content"].append(text)
            elif current_sub:
                current_sub["content"].append(text)
            else:
                child = {"heading": text, "content": [], "children": []}
                current_top["children"].append(child)
                current_sub = child
            continue

        # ------------------------------------------------------------------
        #  Plain text – attach to deepest active node
        # ------------------------------------------------------------------
        if current_subsub:
            current_subsub["content"].append(text)
        elif current_sub:
            current_sub["content"].append(text)
        elif current_top:
            current_top["content"].append(text)

        # ------------------------------------------------------------------
    #  Flush any buffered owners
    # ------------------------------------------------------------------
    if capturing_owner_block and current_top and owner_buffer:
        for line in owner_buffer:
            token = line.strip()
            if token:
                current_top["children"].append(
                    {"heading": token, "content": [], "children": []}
                )
        owner_buffer = []
        capturing_owner_block = False

    # ------------------------------------------------------------------
    #  Flush any buffered designees
    # ------------------------------------------------------------------
    if capturing_designee_block and current_top and designee_buffer:
        for line in designee_buffer:
            token = line.strip()
            if token:
                current_top["children"].append(
                    {"heading": token, "content": [], "children": []}
                )
        designee_buffer = []
        capturing_designee_block = False

    # ------------------------------------------------------------------
    #  Revision history
    # ------------------------------------------------------------------
    rev_hist = extract_revision_history(doc)
    if rev_hist["type"] == "table":
        # Only clear text that belongs to a *Revisions* heading
        def _is_revision(node):
            return node and "revision" in normalize(strip_numeric_prefix(node["heading"]))

        if _is_revision(current_subsub):
            current_subsub["content"] = []
        elif _is_revision(current_sub):
            current_sub["content"] = []
        elif _is_revision(current_top):
            current_top["content"] = []
        return {
            "document_title": document_title,
            "purpose_scope_block": purpose_scope_block,
            "sections": sections,
            "revision_history": rev_hist,
        }

    # Written rev history → attach to last node (cleared afterwards)
    written = {"type": "written", "content": []}
    def _is_revision(node):
        return node and "revision" in normalize(strip_numeric_prefix(node["heading"]))

    if _is_revision(current_subsub):
        written["content"] = current_subsub["content"]
        current_subsub["content"] = []
    elif _is_revision(current_sub):
        written["content"] = current_sub["content"]
        current_sub["content"] = []
    elif _is_revision(current_top):
        written["content"] = current_top["content"]
        current_top["content"] = []

    return {
        "document_title": document_title,
        "purpose_scope_block": purpose_scope_block,
        "sections": sections,
        "revision_history": written,
    }
